"""
generate_docs.py — Job Application Document Generator
Usage: python generate_docs.py <path_to_content.json>

Expects a JSON file with the following schema:
{
  "output_folder": "/path/to/project/Applications/2026-04-16/CompanyName",
  "stack": "Angular",
  "lang": "EN",
  "resume_en": {
    "summary": "...",
    "skills": {
      "frontend": "Angular (2+), TypeScript, ...",
      "tools": "Jest, Git, ...",
      "methodologies": "Agile, ...",
      "languages": "English (Fluent), ..."
    },
    "experience": [
      {
        "title": "Senior Frontend Developer (Angular)",
        "company": "Example Corp (via contractor)",
        "period": "Jun 2025 - March 2026",
        "subtitle": "Enterprise Procurement Platform | USA (Global)",
        "bullets": ["bullet 1", "bullet 2"],
        "stack_line": "Stack: Angular 19, TypeScript, ..."
      }
    ],
    "education": "Example University - Bachelor, Computer Science",
    "courses": "Angular Updates Course, Angular Advanced Course, ..."
  },
  "resume_pl": null,
  "cover_letter_en": "Full cover letter text...",
  "cover_letter_pl": "Pełny tekst listu motywacyjnego...",
  "about_me_en": "3-5 sentence elevator pitch...",
  "about_me_pl": "3-5 zdań elevator pitch..."
}
"""

import json
import os
import re
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from hunter import candidate, gen_profile
from hunter.services.tracker_service import record_successful_apply

# Fallback document-layout defaults, used when generation.yaml/gen_profile
# don't resolve a value — kept as the doc-comment anchor for what each
# document.* key means; see hunter.gen_profile.builtin_defaults() for the
# authoritative copy. Read at call time (not as function-signature defaults)
# so a profile change within the same process takes effect immediately —
# see hunter.filters._resolve_flt's docstring for why.
_DEFAULT_FONT = "Calibri"
_DEFAULT_SIZES = {"name": 16.0, "headline": 13.0, "body": 11.0, "small": 10.0}
_DEFAULT_MARGINS_CM = {"top": 0.8, "bottom": 0.5, "left": 1.0, "right": 1.0}
_DEFAULT_SECTIONS = ["SUMMARY", "SKILLS", "EXPERIENCE", "EDUCATION", "ADDITIONAL COURSES"]
_DEFAULT_SKILL_CATEGORIES = [
    {"key": "frontend", "label": "Frontend"},
    {"key": "tools", "label": "Tools"},
    {"key": "methodologies", "label": "Methodologies"},
    {"key": "languages", "label": "Languages"},
]

# Polish boards (e.g. Pracuj.pl) often reject uploads when the filename exceeds ~50 characters.
MAX_ATTACHMENT_BASENAME_LEN = 50

# GDPR/RODO consent clause. Static legal text — never LLM-generated. Rendered as the
# last paragraph of the resume BODY (not a docx footer) so ATS parsers still read it.
# Controlled by CV_GDPR_CLAUSE in hunter.config ("both" / "pl" / "none").
GDPR_CLAUSE_PL = (
    "Wyrażam zgodę na przetwarzanie danych osobowych zawartych w niniejszym "
    "dokumencie do realizacji obecnego oraz przyszłych procesów rekrutacji zgodnie "
    "z ustawą z dnia 10 maja 2018 roku o ochronie danych osobowych (Dz.U. 2018 poz. "
    "1000) oraz zgodnie z art. 6 ust. 1 lit. a RODO (Rozporządzenie Parlamentu "
    "Europejskiego i Rady (UE) 2016/679 z dnia 27 kwietnia 2016 r.)."
)
GDPR_CLAUSE_EN = (
    "I hereby consent to the processing of my personal data for the purposes "
    "necessary to carry out the current and future recruitment processes, in "
    "accordance with the Act of 10 May 2018 on the Protection of Personal Data "
    "(Journal of Laws 2018, item 1000) and Article 6(1)(a) of Regulation (EU) "
    "2016/679 of the European Parliament and of the Council of 27 April 2016 (GDPR)."
)


def _safe_stack_segment(stack: str, max_len: int = 22) -> str:
    s = (stack or "FE").strip()
    s = re.sub(r"[^\w\-+]", "_", s)
    s = re.sub(r"_+", "_", s).strip("_")
    s = s.replace("+", "_")
    return (s[:max_len] or "FE").strip("_") or "FE"


def resume_docx_basename(stack: str, lang: str) -> str:
    """CV file basename for DOCX/PDF; length <= MAX_ATTACHMENT_BASENAME_LEN (incl. .docx)."""
    lang_u = (lang or "EN").strip().upper()[:2] or "EN"
    ext = ".docx"
    prefix = candidate.get("identity.cv_filename_prefix", candidate.DEFAULT_CV_FILENAME_PREFIX)
    year = datetime.now().year
    # Fixed parts: "{prefix}_" + "_{year}_XX.docx" reserved around the stack segment
    reserved = len(prefix) + 1 + len(f"_{year}_{lang_u}{ext}")
    max_stack = MAX_ATTACHMENT_BASENAME_LEN - reserved
    safe = _safe_stack_segment(stack, max_len=max_stack)
    primary = f"{prefix}_{safe}_{year}_{lang_u}{ext}"
    if len(primary) <= MAX_ATTACHMENT_BASENAME_LEN:
        return primary
    fallback = f"{prefix}_{year}_{lang_u}{ext}"
    if len(fallback) <= MAX_ATTACHMENT_BASENAME_LEN:
        return fallback
    return f"CV_{year}_{lang_u}{ext}"


def set_font(run, name=None, size=11, bold=False, italic=False, color=None):
    """name=None resolves document.font (generation.yaml) at call time, not
    as a frozen function-signature default — every call site below omits
    `name` except where a document element intentionally uses a fixed font."""
    if name is None:
        name = gen_profile.get("document.font", _DEFAULT_FONT)
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_horizontal_line(paragraph):
    """Add a thin bottom border to a paragraph (acts as a divider)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)


def keep_with_next(paragraph, keep_lines=True):
    """Prevent the paragraph from being split from what follows across a page break.

    w:keepNext  -> stay on the same page as the next paragraph (no orphaned heading)
    w:keepLines -> do not break the paragraph's own lines across pages
    """
    pPr = paragraph._p.get_or_add_pPr()
    keep_next = OxmlElement("w:keepNext")
    pPr.append(keep_next)
    if keep_lines:
        keep_lines_el = OxmlElement("w:keepLines")
        pPr.append(keep_lines_el)


def set_paragraph_spacing(paragraph, before=0, after=4, line_spacing=1.08):
    from docx.oxml.ns import qn

    pPr = paragraph._p.get_or_add_pPr()
    pSpacing = OxmlElement("w:spacing")
    pSpacing.set(qn("w:before"), str(int(before * 20)))
    pSpacing.set(qn("w:after"), str(int(after * 20)))
    pSpacing.set(qn("w:line"), str(int(line_spacing * 240)))
    pSpacing.set(qn("w:lineRule"), "auto")
    pPr.append(pSpacing)


def add_section_heading(doc, text):
    sizes = gen_profile.get("document.sizes", _DEFAULT_SIZES)
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    set_font(run, size=sizes.get("body", _DEFAULT_SIZES["body"]), bold=True)
    add_horizontal_line(p)
    set_paragraph_spacing(p, before=6, after=3)
    keep_with_next(p)
    return p


def add_gdpr_clause(doc, lang):
    """Append the GDPR/RODO consent clause as the last body paragraph.

    Small italic grey text. Gated by CV_GDPR_CLAUSE ("both" / "pl" / "none").
    lang is "EN" or "PL"; EN CVs are skipped unless mode is "both".
    """
    try:
        from hunter.config import CV_GDPR_CLAUSE
    except Exception:
        CV_GDPR_CLAUSE = "both"
    mode = (CV_GDPR_CLAUSE or "both").strip().lower()
    lang_u = (lang or "EN").strip().upper()[:2]
    if mode == "none":
        return
    if mode == "pl" and lang_u != "PL":
        return
    text = GDPR_CLAUSE_PL if lang_u == "PL" else GDPR_CLAUSE_EN
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=7.5, italic=True, color=(120, 120, 120))
    set_paragraph_spacing(p, before=8, after=0)


def build_resume(doc, data, stack, lang="EN"):
    sizes = gen_profile.get("document.sizes", _DEFAULT_SIZES)
    name_size = sizes.get("name", _DEFAULT_SIZES["name"])
    headline_size = sizes.get("headline", _DEFAULT_SIZES["headline"])
    body_size = sizes.get("body", _DEFAULT_SIZES["body"])
    small_size = sizes.get("small", _DEFAULT_SIZES["small"])
    # Heading LABELS by fixed structural position (0=summary .. 4=courses) —
    # only the text is configurable, not which content renders where. See
    # hunter.gen_profile.builtin_defaults()'s "document.sections" comment for
    # why position 2 ("EXPERIENCE") is an ATS-parsing risk to rename.
    sections = gen_profile.get("document.sections", _DEFAULT_SECTIONS)

    def _heading(position: int) -> str:
        if isinstance(sections, list) and position < len(sections) and sections[position]:
            return sections[position]
        return _DEFAULT_SECTIONS[position]

    name = candidate.get("identity.full_name", candidate.DEFAULT_FULL_NAME)
    aka = candidate.get("identity.aka", candidate.DEFAULT_AKA)
    subtitle = f"also known as {aka}" if aka else ""
    headline_base = candidate.get("identity.headline", candidate.DEFAULT_HEADLINE)
    headline = f"{headline_base} ({stack})"
    contact = candidate.get("identity.contact", candidate.DEFAULT_CONTACT)

    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(name)
    set_font(run, size=name_size, bold=True)
    set_paragraph_spacing(p, before=0, after=2)

    # Subtitle (also known as) — optional, skipped when aka is blank
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle)
        set_font(run, size=small_size, italic=True)
        set_paragraph_spacing(p, before=0, after=2)

    # Headline
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(headline)
    set_font(run, size=headline_size, bold=True)
    set_paragraph_spacing(p, before=0, after=2)

    # Contact
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(contact)
    set_font(run, size=small_size)
    set_paragraph_spacing(p, before=0, after=6)

    # SUMMARY
    add_section_heading(doc, _heading(0))
    p = doc.add_paragraph()
    run = p.add_run(data["summary"])
    set_font(run, size=body_size)
    set_paragraph_spacing(p, before=3, after=6)

    # SKILLS
    add_section_heading(doc, _heading(1))
    skills = data["skills"]
    categories = gen_profile.get("document.skill_categories", _DEFAULT_SKILL_CATEGORIES)
    skill_lines = [
        (c.get("label", ""), skills.get(c.get("key", ""), ""))
        for c in categories
        if isinstance(c, dict)
    ]
    for label, value in skill_lines:
        if value:
            p = doc.add_paragraph()
            run_label = p.add_run(f"{label}: ")
            set_font(run_label, size=body_size, bold=True)
            run_value = p.add_run(value)
            set_font(run_value, size=body_size)
            set_paragraph_spacing(p, before=1, after=1)

    # EXPERIENCE — Taleo and a handful of legacy ATS expect an exact-match
    # section header. "WORK EXPERIENCE" can be classified as "Other" and
    # the entire experience array silently drops out of the parsed fields.
    add_section_heading(doc, _heading(2))
    # Date normalizer (lazy import — generate_docs.py is invoked as a script).
    from hunter.date_normalize import normalize_period

    for job in data["experience"]:
        # Title | Company — Period (period normalized to MM/YYYY – MM/YYYY for
        # Taleo-strict ATS parsers; unparseable periods pass through unchanged).
        p = doc.add_paragraph()
        run_title = p.add_run(f"{job['title']} | {job['company']}")
        set_font(run_title, size=body_size, bold=True)
        run_period = p.add_run(f"   {normalize_period(job['period'])}")
        set_font(run_period, size=small_size, italic=True)
        set_paragraph_spacing(p, before=4, after=1)
        keep_with_next(p)

        # Subtitle / context line
        if job.get("subtitle"):
            p = doc.add_paragraph()
            run = p.add_run(job["subtitle"])
            set_font(run, size=small_size, italic=True)
            set_paragraph_spacing(p, before=0, after=2)
            keep_with_next(p)

        # Bullets
        for bullet in job.get("bullets", []):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(bullet)
            set_font(run, size=body_size)
            set_paragraph_spacing(p, before=0, after=1)

        # Stack line
        if job.get("stack_line"):
            p = doc.add_paragraph()
            run = p.add_run(job["stack_line"])
            set_font(run, size=small_size, bold=True)
            set_paragraph_spacing(p, before=1, after=3)

    # EDUCATION
    add_section_heading(doc, _heading(3))
    p = doc.add_paragraph()
    edu = data.get("education", "")
    if not isinstance(edu, str):
        edu = (
            ", ".join(str(v) for v in edu.values() if v)
            if isinstance(edu, dict)
            else str(edu or "")
        )
    run = p.add_run(edu)
    set_font(run, size=body_size)
    set_paragraph_spacing(p, before=3, after=3)

    # ADDITIONAL COURSES
    add_section_heading(doc, _heading(4))
    p = doc.add_paragraph()
    courses = data.get("courses", "")
    if not isinstance(courses, str):
        courses = (
            ", ".join(str(v) for v in courses) if isinstance(courses, list) else str(courses or "")
        )
    run = p.add_run(courses)
    set_font(run, size=body_size)
    set_paragraph_spacing(p, before=3, after=3)

    # GDPR/RODO consent clause — last body paragraph (ATS-readable, not a footer)
    add_gdpr_clause(doc, lang)


def build_cover_letter(doc, text):
    body_size = gen_profile.get("document.sizes", _DEFAULT_SIZES).get(
        "body", _DEFAULT_SIZES["body"]
    )
    for line in text.split("\n"):
        p = doc.add_paragraph()
        run = p.add_run(line)
        set_font(run, size=body_size)
        set_paragraph_spacing(p, before=0, after=6)


def set_margins(doc, top_cm=None, bottom_cm=None, left_cm=None, right_cm=None):
    """Any margin left None resolves document.margins_cm (generation.yaml) at
    call time — callers that pass no arguments (the normal case) get the
    profile-driven default; an explicit value still wins."""
    margins = gen_profile.get("document.margins_cm", _DEFAULT_MARGINS_CM)
    if top_cm is None:
        top_cm = margins.get("top", _DEFAULT_MARGINS_CM["top"])
    if bottom_cm is None:
        bottom_cm = margins.get("bottom", _DEFAULT_MARGINS_CM["bottom"])
    if left_cm is None:
        left_cm = margins.get("left", _DEFAULT_MARGINS_CM["left"])
    if right_cm is None:
        right_cm = margins.get("right", _DEFAULT_MARGINS_CM["right"])
    for section in doc.sections:
        section.top_margin = Cm(top_cm)
        section.bottom_margin = Cm(bottom_cm)
        section.left_margin = Cm(left_cm)
        section.right_margin = Cm(right_cm)


def set_author(doc, name=None):
    if name is None:
        name = candidate.get("identity.full_name", candidate.DEFAULT_FULL_NAME)
    props = doc.core_properties
    props.author = name
    props.last_modified_by = name
    props.title = ""
    props.subject = ""
    props.keywords = ""


def save_docx(doc, path_docx):
    """Save DOCX only. PDFs are converted in bulk at the end."""
    set_author(doc)
    doc.save(path_docx)
    print(f"  [OK] DOCX: {path_docx}")


def convert_all_to_pdf(output_folder):
    """Convert all DOCX files in the folder to PDF in a single LibreOffice call."""
    import subprocess
    from hunter.config import SOFFICE_PATH

    soffice = SOFFICE_PATH
    docx_files = list(Path(output_folder).glob("*.docx"))
    if not docx_files:
        return
    result = subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(output_folder)]
        + [str(f) for f in docx_files],
        capture_output=True,
        text=True,
        timeout=120,
    )
    if result.returncode == 0:
        for f in docx_files:
            print(f"  [OK] PDF:  {str(f).replace('.docx', '.pdf')}")
    else:
        print(f"  [WARN] PDF conversion failed: {result.stderr.strip()}")


def update_tracker(content: dict, force_mode: bool = False) -> None:
    """Write successful apply record through centralized tracker API."""
    written = record_successful_apply(content, force=force_mode)
    if written:
        print("  [OK] Tracker updated")
    else:
        apply_url = content.get("apply_url", "")
        print(f"  [tracker] Skipping — successful entry already exists for {str(apply_url)[:60]}")


def main():
    if len(sys.argv) < 2:
        print("Usage: python generate_docs.py <content.json> [--full] [--force] [--no-tracker]")
        sys.exit(1)

    # Identity gate — before any file is written. Without a configured
    # candidate.yaml every identity field falls back to a placeholder, and a
    # rendered PDF is the one artifact that leaves the machine and reaches a
    # real employer. Aborting here is a normal retryable failure; a CV sent
    # under the wrong name is not.
    try:
        candidate.require_identity()
    except candidate.CandidateIdentityMissing as exc:
        print(f"[generate_docs] ERROR: {exc}")
        sys.exit(1)

    json_path = sys.argv[1]
    full_mode = "--full" in sys.argv
    force_mode = "--force" in sys.argv
    # --no-tracker: render docs only, skip the tracker write. Used by the
    # dual-apply shadow run so the comparison CV never lands a tracker row.
    no_tracker = "--no-tracker" in sys.argv

    with open(json_path, "r", encoding="utf-8") as f:
        content = json.load(f)

    # Belt-and-suspenders: sanitize company names + education/courses even if
    # apply_agent already ran sanitizer (handles manual / re-run scenarios).
    try:
        from hunter.resume_sanitizer import sanitize_content

        content = sanitize_content(content)
        # Persist corrected content back to disk
        Path(json_path).write_text(
            json.dumps(content, ensure_ascii=False, indent=2), encoding="utf-8"
        )
    except Exception as _san_err:
        print(f"  [WARN] resume sanitizer: {_san_err}")

    output_folder = content["output_folder"]
    stack = content["stack"]

    os.makedirs(output_folder, exist_ok=True)
    mode_label = "FULL" if full_mode else "SHORT (PDF-only, EN CV)"
    print(f"\nOutput folder: {output_folder}")
    print(f"Mode: {mode_label}\n")

    # --- Resume EN ---
    if content.get("resume_en"):
        doc = Document()
        set_margins(doc)
        build_resume(doc, content["resume_en"], stack, "EN")
        fname = resume_docx_basename(stack, "EN")
        save_docx(doc, Path(output_folder) / fname)

    # --- Resume PL ---
    # Rendered in full mode, OR in short mode when the posting is Polish — a Polish
    # employer should receive the (clean) Polish CV as the primary document, so it
    # ships even in the default short flow alongside the EN CV.
    _primary_pl = (content.get("primary_lang") or "").strip().upper() == "PL"
    if (full_mode or _primary_pl) and content.get("resume_pl"):
        doc = Document()
        set_margins(doc)
        build_resume(doc, content["resume_pl"], stack, "PL")
        fname = resume_docx_basename(stack, "PL")
        save_docx(doc, Path(output_folder) / fname)

    # --- Cover Letter EN ---
    if content.get("cover_letter_en"):
        doc = Document()
        set_margins(doc)
        build_cover_letter(doc, content["cover_letter_en"])
        save_docx(doc, Path(output_folder) / "Cover_Letter_EN.docx")

    # --- Cover Letter PL ---
    if content.get("cover_letter_pl"):
        doc = Document()
        set_margins(doc)
        build_cover_letter(doc, content["cover_letter_pl"])
        save_docx(doc, Path(output_folder) / "Cover_Letter_PL.docx")

    # --- About Me ---
    _about_me_pl = os.getenv("GENERATE_ABOUT_ME_PL", "true").lower() in ("true", "1", "yes")
    if full_mode or _about_me_pl:
        from hunter.about_me_agent import generate_about_me

        if full_mode:
            generate_about_me(Path(output_folder), lang="en")
        generate_about_me(Path(output_folder), lang="pl")

    # --- Update tracker.xlsx before PDF step so a LibreOffice crash doesn't lose the record ---
    if no_tracker:
        print("  [tracker] Skipping (--no-tracker: shadow/comparison run)")
    else:
        try:
            update_tracker(content, force_mode=force_mode)
        except Exception as exc:
            print(f"  [WARN] Tracker update failed (close tracker.xlsx if open): {exc}")

    # --- Convert all DOCX to PDF in one shot ---
    try:
        convert_all_to_pdf(output_folder)
    except Exception as exc:
        print(f"  [WARN] PDF conversion failed: {exc}")

    # --- Short mode: remove DOCX only when a sibling PDF exists (LibreOffice OK) ---
    if not full_mode:
        out = Path(output_folder)
        for docx_file in out.glob("*.docx"):
            pdf_candidate = out / (docx_file.stem + ".pdf")
            if pdf_candidate.is_file():
                docx_file.unlink()
                print(f"  [cleanup] Removed intermediate: {docx_file.name}")
            else:
                print(
                    f"  [cleanup] Kept DOCX (no PDF yet): {docx_file.name} "
                    "- check LibreOffice / run convert manually"
                )

    print("\nDone!\n")


if __name__ == "__main__":
    main()
