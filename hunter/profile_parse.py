"""hunter/profile_parse.py — resume text extraction + LLM parsing into a
structured Profile (hunter/profile_schema.py).

docs/RESUME_PROFILE_STORE_PLAN.md M3. Two layers:

  1. `extract_resume_text(path)` — read raw text out of an uploaded .docx /
     .pdf / .txt / .md file. This step CAN fail (a corrupt file, an unknown
     extension) and raises `ProfileParseError` when it does — there is no
     text to hand the rest of the pipeline.
  2. `parse_resume_text(text, llm=...)` — turn that text into a Profile.
     This step must NEVER hard-fail: a resume upload is an onboarding
     event, and a bad LLM call/response/validation result degrades to the
     same fallback `parse_resume_text(text)` produces with no `llm` at all
     (the whole text as one leftover) rather than raising or losing data.

The `llm` parameter is a plain injected callable matching
`llm_client.call_llm`'s keyword signature (system_prompt, user_message,
provider, model, api_key, ...) -> dict — the same dependency-injection shape
`hunter/prescreen.py` uses internally, except here the caller supplies it
(so `tools/parse_resume.py --no-llm` can simply pass `llm=None` for a $0
run, and tests can inject a stub without touching a real model).
"""

from __future__ import annotations

import json
import logging
import re
from pathlib import Path
from typing import Any, Callable

from hunter import contact_extract, profile_schema
from hunter.config import JUDGE_API_KEY, JUDGE_MODEL, JUDGE_PROVIDER
from hunter.pipeline.folders import PROMPTS_DIR

logger = logging.getLogger(__name__)

_TEXT_EXTENSIONS = {".txt", ".md"}

LLMCallable = Callable[..., dict]

_PROMPT_PATH = PROMPTS_DIR / "resume_parse.md"
_MAX_INPUT_CHARS = 20000


class ProfileParseError(RuntimeError):
    """Raised when a resume file cannot be turned into text at all — unknown
    extension, corrupt/unreadable file, or a file with no extractable text."""


def _extract_docx_text(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n".join(parts)


def _extract_pdf_text(path: Path) -> str:
    # Reuse the same extractor hunter/ats_pdf_roundtrip.py uses to read a
    # rendered CV PDF back — one PDF->text implementation for the whole repo.
    from hunter.pdf_text import extract_pdf_text

    return extract_pdf_text(path)


def extract_resume_text(path: Path) -> str:
    """Extract plain text from an uploaded resume file.

    Supports `.docx` (python-docx: paragraphs + table cells), `.pdf`
    (`hunter.pdf_text.extract_pdf_text`), and `.txt`/`.md` (read as-is).
    Any other extension, an unreadable file, or a file with no extractable
    text raises `ProfileParseError` with a message safe to show the user who
    uploaded it.
    """
    path = Path(path)
    suffix = path.suffix.lower()

    try:
        if suffix == ".docx":
            text = _extract_docx_text(path)
        elif suffix == ".pdf":
            text = _extract_pdf_text(path)
        elif suffix in _TEXT_EXTENSIONS:
            text = path.read_text(encoding="utf-8", errors="replace")
        else:
            raise ProfileParseError(f"Unsupported resume file type: {suffix or '(none)'}")
    except ProfileParseError:
        raise
    except Exception as e:
        raise ProfileParseError(f"Could not read {path.name}: {e}") from e

    if not text.strip():
        raise ProfileParseError(f"No extractable text in {path.name}")
    return text


# ── Text -> Profile ─────────────────────────────────────────────────────────


def _contact_line(text: str) -> str:
    """Best-effort "phone | email" line, deterministic, $0.

    `contact_extract.extract_contacts` was built for finding a RECRUITER's
    contact inside a job posting, but its email/phone regexes are generic —
    over a resume's own header (no "recruiter"/"kontakt" label to match) it
    still finds the candidate's own email and, if present, attaches the
    first phone number to it. Good enough for a pre-fill the user reviews
    on the confirmation screen; the candidate's own name is deliberately
    NOT guessed from this (see parse_resume_text's docstring)."""
    contacts = contact_extract.extract_contacts(text)
    if not contacts:
        return ""
    parts = [p for p in (contacts[0].phone, contacts[0].email) if p]
    return " | ".join(parts)


def parse_resume_text(
    text: str,
    llm: LLMCallable | None = None,
    *,
    source_upload_id: str = "",
) -> profile_schema.Profile:
    """Turn resume text into a structured Profile.

    With `llm=None`, no model call is made at all: the whole text lands in
    leftovers plus a deterministic email/phone pre-fill (this is the `$0
    --no-llm` mode). With an `llm` callable, one cheap model call attempts a
    real parse; any call failure, malformed response, or failed
    `profile_schema.validate()` degrades to that exact same fallback — the
    parse never hard-fails and never returns half a document silently
    passed off as fully structured. The candidate's own NAME is never
    guessed by the fallback path; only a validated LLM response may set it,
    and the user still confirms it on the editor's confirmation screen.
    """
    text = (text or "").strip()
    if llm is not None and text:
        parsed = _try_llm_parse(text, llm, source_upload_id=source_upload_id)
        if parsed is not None:
            return parsed
    return _fallback_profile(text, source_upload_id=source_upload_id)


def _fallback_profile(text: str, *, source_upload_id: str = "") -> profile_schema.Profile:
    """The parse-never-fails branch: no structure extracted, just a contact
    pre-fill and the raw text preserved as a single leftover for the user to
    reassign by hand."""
    profile = profile_schema.Profile()
    if not text:
        return profile
    contact = _contact_line(text)
    if contact:
        profile.core.identity.contact = contact
    profile.leftovers = [profile_schema.Leftover(text=text, source_upload_id=source_upload_id)]
    return profile


# ── LLM-assisted parse ───────────────────────────────────────────────────────


def _system_prompt() -> str:
    return _PROMPT_PATH.read_text(encoding="utf-8")


def _user_message(text: str) -> str:
    return "--- RESUME TEXT ---\n" + text[:_MAX_INPUT_CHARS]


def _normalize_leftovers(raw: Any) -> list[dict]:
    """Accept either the documented `[{"text": "..."}]` shape or a plain
    `["...", "..."]` list of strings — models drift toward the simpler shape
    despite the prompt, and a leftover is exactly the kind of content we
    must not silently drop for being shaped slightly wrong."""
    if not isinstance(raw, list):
        return []
    out: list[dict] = []
    for item in raw:
        if isinstance(item, str):
            if item.strip():
                out.append({"text": item})
        elif isinstance(item, dict) and str(item.get("text", "")).strip():
            out.append(item)
    return out


def _profile_dict_from_llm_response(raw: Any) -> dict | None:
    """Tolerant reshape of the model's raw JSON into the
    `{"schema_version", "core", "leftovers"}` document `profile_schema.from_dict`
    expects. Returns None for anything that isn't at least a dict with a
    `core` object — that's the minimum a caller can trust."""
    if isinstance(raw, str):
        try:
            raw = json.loads(raw)
        except (TypeError, ValueError):
            return None
    if not isinstance(raw, dict):
        return None
    core = raw.get("core")
    if not isinstance(core, dict):
        return None
    return {
        "schema_version": profile_schema.SCHEMA_VERSION,
        "core": core,
        "leftovers": _normalize_leftovers(raw.get("leftovers")),
    }


def _ensure_cv_filename_prefix(profile: profile_schema.Profile) -> None:
    """Derive a filename-safe slug from the extracted name when the model
    didn't set one — this is a rendering detail, not a fact about the
    candidate, so deriving it deterministically isn't "guessing a fact"."""
    identity = profile.core.identity
    if identity.cv_filename_prefix.strip() or not identity.full_name.strip():
        return
    slug = re.sub(r"[^A-Za-z0-9]+", "_", identity.full_name.strip()).strip("_")
    if slug:
        identity.cv_filename_prefix = f"{slug}_CV"


def _try_llm_parse(
    text: str, llm: LLMCallable, *, source_upload_id: str
) -> profile_schema.Profile | None:
    """One cheap-model parse attempt. Returns None — never raises — on any
    call failure, malformed response, or failed identity validation, so the
    caller can fall back to the deterministic leftovers-only branch."""
    try:
        raw = llm(
            system_prompt=_system_prompt(),
            user_message=_user_message(text),
            provider=JUDGE_PROVIDER,
            model=JUDGE_MODEL,
            api_key=JUDGE_API_KEY,
            max_tokens=4096,
        )
    except Exception as e:  # noqa: BLE001 — logged here, caller falls back
        logger.warning("profile_parse: LLM call failed, falling back to leftovers: %s", e)
        return None

    doc = _profile_dict_from_llm_response(raw)
    if doc is None:
        logger.warning("profile_parse: LLM response malformed, falling back to leftovers")
        return None

    profile = profile_schema.from_dict(doc)
    for leftover in profile.leftovers:
        if not leftover.source_upload_id:
            leftover.source_upload_id = source_upload_id

    _ensure_cv_filename_prefix(profile)
    if not profile.core.identity.contact.strip():
        contact = _contact_line(text)
        if contact:
            profile.core.identity.contact = contact

    problems = profile_schema.validate(profile)
    if problems:
        logger.warning("profile_parse: LLM response failed validation: %s", problems)
        return None
    return profile
