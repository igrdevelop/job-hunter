"""Document-layout parity + configurability for generate_docs.py (wave 3
PR 2, docs/GENERATION_ARCHITECTURE_ANALYSIS.md §6).

1. build_resume()/build_cover_letter()/set_margins() with builtin defaults
   (no candidate/generation.yaml present) reproduce EXACTLY the paragraph
   sequence (text, font, size, bold, italic) and margins captured from
   generate_docs.py BEFORE this wave — pinned in
   tests/fixtures/generate_docs_resume_golden.json, generated from the
   pre-PR2 code (git history) with the same fixture data used here.
2. A generation.yaml override (font, sizes, margins, section labels, skill
   categories, GDPR mode) actually changes the rendered output — proves the
   knobs are real, not just documented.

PDF bytes are never compared (timestamps) — only python-docx's in-memory
paragraph/run structure, which is what an ATS or a human actually reads.
"""

from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document

import generate_docs
from hunter import candidate, gen_profile

FIXTURE_DIR = Path(__file__).resolve().parent / "fixtures"
GOLDEN_PATH = FIXTURE_DIR / "generate_docs_resume_golden.json"

CANDIDATE_YAML = """
identity:
  full_name: "Jane Golden Tester"
  aka: "JGT"
  headline: "Senior Test Engineer"
  contact: "jane@example.com | +1 555 0100 | Testville"
  cv_filename_prefix: "Golden_CV"
"""

FIXTURE_DATA = {
    "summary": "A concise professional summary for the golden snapshot test.",
    "skills": {
        "frontend": "Angular, TypeScript, RxJS",
        "tools": "Git, Jest, Webpack",
        "methodologies": "Agile, Scrum, TDD",
        "languages": "English (Fluent), Polish (B2)",
    },
    "experience": [
        {
            "title": "Senior Frontend Developer",
            "company": "Acme Corp",
            "period": "Jan 2020 - Present",
            "subtitle": "Enterprise Platform | Remote",
            "bullets": ["Built feature X", "Improved performance by 30%"],
            "stack_line": "Stack: Angular 17, TypeScript, RxJS",
        },
        {
            "title": "Frontend Developer",
            "company": "Beta LLC",
            "period": "Jun 2017 - Dec 2019",
            "bullets": ["Shipped feature Y"],
        },
    ],
    "education": "Golden University - BSc Computer Science",
    "courses": "Angular Advanced, TypeScript Deep Dive",
}

COVER_LETTER_TEXT = "Dear Hiring Manager,\n\nI am excited to apply.\n\nSincerely,\nJane"


@pytest.fixture
def candidate_stub(tmp_path):
    cand_path = tmp_path / "candidate.yaml"
    cand_path.write_text(CANDIDATE_YAML, encoding="utf-8")
    candidate._set_path(cand_path)
    yield
    candidate._set_path(None)


@pytest.fixture(autouse=True)
def _clean_gen_profile_cache(monkeypatch):
    monkeypatch.delenv("GENERATION_YAML_PATH", raising=False)
    monkeypatch.delenv("CV_GDPR_CLAUSE", raising=False)
    gen_profile.clear_gen_profile_cache()
    yield
    gen_profile.clear_gen_profile_cache()


def _snapshot_doc(doc) -> list[dict]:
    out = []
    for p in doc.paragraphs:
        runs = []
        for r in p.runs:
            runs.append(
                {
                    "text": r.text,
                    "font_name": r.font.name,
                    "size_pt": r.font.size.pt if r.font.size else None,
                    "bold": r.font.bold,
                    "italic": r.font.italic,
                }
            )
        out.append({"alignment": str(p.alignment) if p.alignment else None, "runs": runs})
    return out


def _build_resume_doc():
    doc = Document()
    generate_docs.set_margins(doc)
    generate_docs.build_resume(doc, FIXTURE_DATA, "Angular", "EN")
    return doc


def _build_cover_doc():
    doc = Document()
    generate_docs.set_margins(doc)
    generate_docs.build_cover_letter(doc, COVER_LETTER_TEXT)
    return doc


def test_default_resume_matches_golden_snapshot(candidate_stub, tmp_path):
    """No generation.yaml present => byte-for-byte the pre-wave-3 output."""
    golden = json.loads(GOLDEN_PATH.read_text(encoding="utf-8"))
    doc = _build_resume_doc()
    assert _snapshot_doc(doc) == golden["resume"]
    margins = doc.sections[0]
    assert margins.top_margin.cm == pytest.approx(golden["margins_cm"]["top"], rel=1e-3)
    assert margins.bottom_margin.cm == pytest.approx(golden["margins_cm"]["bottom"], rel=1e-3)
    assert margins.left_margin.cm == pytest.approx(golden["margins_cm"]["left"], rel=1e-3)
    assert margins.right_margin.cm == pytest.approx(golden["margins_cm"]["right"], rel=1e-3)


def test_default_cover_letter_matches_golden_snapshot(candidate_stub):
    golden = json.loads(GOLDEN_PATH.read_text(encoding="utf-8"))
    doc = _build_cover_doc()
    assert _snapshot_doc(doc) == golden["cover_letter"]


def test_font_override_applies(candidate_stub, tmp_path, monkeypatch):
    path = tmp_path / "generation.yaml"
    path.write_text("document:\n  font: Arial\n", encoding="utf-8")
    monkeypatch.setenv("GENERATION_YAML_PATH", str(path))
    doc = _build_resume_doc()
    assert doc.paragraphs[0].runs[0].font.name == "Arial"


def test_sizes_override_applies(candidate_stub, tmp_path, monkeypatch):
    path = tmp_path / "generation.yaml"
    path.write_text("document:\n  sizes: {name: 20, body: 12}\n", encoding="utf-8")
    monkeypatch.setenv("GENERATION_YAML_PATH", str(path))
    doc = _build_resume_doc()
    assert doc.paragraphs[0].runs[0].font.size.pt == 20.0
    # headline/small untouched by the partial override
    headline_run = next(p for p in doc.paragraphs if "Senior Test Engineer" in p.runs[0].text)
    assert headline_run.runs[0].font.size.pt == 13.0


def test_margins_override_applies(candidate_stub, tmp_path, monkeypatch):
    path = tmp_path / "generation.yaml"
    path.write_text("document:\n  margins_cm: {top: 2.0}\n", encoding="utf-8")
    monkeypatch.setenv("GENERATION_YAML_PATH", str(path))
    doc = _build_resume_doc()
    assert doc.sections[0].top_margin.cm == pytest.approx(2.0, rel=1e-3)
    # untouched margins keep the builtin default
    assert doc.sections[0].left_margin.cm == pytest.approx(1.0, rel=1e-3)


def test_section_label_override_applies(candidate_stub, tmp_path, monkeypatch):
    path = tmp_path / "generation.yaml"
    path.write_text(
        'document:\n  sections: ["OVERVIEW", "SKILLS", "EXPERIENCE", "EDUCATION", "COURSES"]\n',
        encoding="utf-8",
    )
    monkeypatch.setenv("GENERATION_YAML_PATH", str(path))
    doc = _build_resume_doc()
    headings = [p.runs[0].text for p in doc.paragraphs if p.runs and p.runs[0].text == "OVERVIEW"]
    assert headings == ["OVERVIEW"]
    assert not any(p.runs and p.runs[0].text == "SUMMARY" for p in doc.paragraphs)


def test_skill_categories_override_applies(candidate_stub, tmp_path, monkeypatch):
    path = tmp_path / "generation.yaml"
    path.write_text(
        "document:\n"
        "  skill_categories:\n"
        "    - {key: frontend, label: Frameworks}\n"
        "    - {key: languages, label: Spoken Languages}\n",
        encoding="utf-8",
    )
    monkeypatch.setenv("GENERATION_YAML_PATH", str(path))
    doc = _build_resume_doc()
    texts = [p.runs[0].text for p in doc.paragraphs if p.runs]
    assert any(t.startswith("Frameworks: ") for t in texts)
    assert any(t.startswith("Spoken Languages: ") for t in texts)
    assert not any(t.startswith("Tools: ") for t in texts)


def test_gdpr_clause_none_removes_clause(candidate_stub, tmp_path, monkeypatch):
    path = tmp_path / "generation.yaml"
    path.write_text("document:\n  gdpr_clause: none\n", encoding="utf-8")
    monkeypatch.setenv("GENERATION_YAML_PATH", str(path))
    import importlib

    config = importlib.reload(importlib.import_module("hunter.config"))
    try:
        doc = _build_resume_doc()
        texts = " ".join(p.runs[0].text for p in doc.paragraphs if p.runs)
        assert "RODO" not in texts and "GDPR" not in texts
    finally:
        importlib.reload(config)


def test_cv_gdpr_clause_env_still_wins_over_yaml(candidate_stub, tmp_path, monkeypatch):
    path = tmp_path / "generation.yaml"
    path.write_text("document:\n  gdpr_clause: none\n", encoding="utf-8")
    monkeypatch.setenv("GENERATION_YAML_PATH", str(path))
    monkeypatch.setenv("CV_GDPR_CLAUSE", "both")
    import importlib

    config = importlib.reload(importlib.import_module("hunter.config"))
    try:
        assert config.CV_GDPR_CLAUSE == "both"
    finally:
        importlib.reload(config)
