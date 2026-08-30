"""Tests for hunter/profile_parse.py.

Step 3a: extract_resume_text (docx/pdf/txt -> plain text, raises
ProfileParseError on anything unreadable).
Step 3b: parse_resume_text's no-LLM fallback branch — the whole text lands
in leftovers, plus a deterministic contact pre-fill; the parse itself never
raises.
Step 3c: parse_resume_text's LLM-assisted branch — a valid response
produces real structure; a malformed response, a raising call, or a
response that fails profile_schema.validate() all fall back to the exact
same leftovers-only branch as step 3b.
"""

from __future__ import annotations

from pathlib import Path

import pytest

from hunter.profile_parse import ProfileParseError, extract_resume_text, parse_resume_text
from hunter.profile_schema import from_dict, to_dict

FIXTURES = Path(__file__).parent / "fixtures" / "resumes"
JANE_TEXT = (FIXTURES / "jane.txt").read_text(encoding="utf-8")


def test_extract_text_from_txt() -> None:
    text = extract_resume_text(FIXTURES / "jane.txt")
    assert "Jane Doe" in text
    assert "jane.doe@example.com" in text


def test_extract_text_from_md(tmp_path: Path) -> None:
    path = tmp_path / "resume.md"
    path.write_text("# Jane Doe\n\nSenior Frontend Developer\n", encoding="utf-8")
    text = extract_resume_text(path)
    assert "Jane Doe" in text


def test_extract_text_from_docx(tmp_path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("Senior Frontend Developer")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Acme Corp"
    table.rows[0].cells[1].text = "Jan 2024 - Present"
    path = tmp_path / "resume.docx"
    doc.save(str(path))

    text = extract_resume_text(path)
    assert "Jane Doe" in text
    assert "Acme Corp" in text
    assert "Jan 2024 - Present" in text


def test_extract_text_unsupported_extension(tmp_path: Path) -> None:
    path = tmp_path / "resume.xyz"
    path.write_text("whatever", encoding="utf-8")
    with pytest.raises(ProfileParseError):
        extract_resume_text(path)


def test_extract_text_unreadable_pdf(tmp_path: Path) -> None:
    path = tmp_path / "broken.pdf"
    path.write_bytes(b"not a real pdf")
    with pytest.raises(ProfileParseError):
        extract_resume_text(path)


def test_extract_text_missing_file(tmp_path: Path) -> None:
    with pytest.raises(ProfileParseError):
        extract_resume_text(tmp_path / "missing.docx")


def test_extract_text_empty_txt_raises(tmp_path: Path) -> None:
    path = tmp_path / "empty.txt"
    path.write_text("   \n\n", encoding="utf-8")
    with pytest.raises(ProfileParseError):
        extract_resume_text(path)


# ── parse_resume_text — no-LLM fallback (step 3b) ────────────────────────────


def test_parse_without_llm_puts_whole_text_in_leftovers() -> None:
    profile = parse_resume_text(JANE_TEXT)
    assert len(profile.leftovers) == 1
    assert profile.leftovers[0].text == JANE_TEXT.strip()
    assert profile.core.roles == []
    assert profile.core.skills == []


def test_parse_without_llm_fills_contact_from_text() -> None:
    profile = parse_resume_text(JANE_TEXT)
    assert "jane.doe@example.com" in profile.core.identity.contact
    assert "+1 555 000 1234" in profile.core.identity.contact
    # The candidate's own name is never guessed from free text.
    assert profile.core.identity.full_name == ""


def test_parse_without_llm_stamps_source_upload_id() -> None:
    profile = parse_resume_text(JANE_TEXT, source_upload_id="upload-1")
    assert profile.leftovers[0].source_upload_id == "upload-1"


def test_parse_without_llm_empty_text_is_empty_profile() -> None:
    profile = parse_resume_text("   ")
    assert profile.leftovers == []
    assert profile.core.identity.contact == ""


def test_parse_without_llm_round_trips() -> None:
    profile = parse_resume_text(JANE_TEXT)
    rebuilt = from_dict(to_dict(profile))
    assert rebuilt.core.identity.contact == profile.core.identity.contact
    assert rebuilt.leftovers[0].text == profile.leftovers[0].text


# ── parse_resume_text — LLM-assisted parse (step 3c) ─────────────────────────


def test_parse_with_llm_valid_response_extracts_structure(fake_llm) -> None:
    fake_llm.judge_response = {
        "core": {
            "identity": {"full_name": "Jane Doe", "headline": "Senior Frontend Developer"},
            "roles": [
                {
                    "company": "Acme Corp",
                    "title": "Senior Frontend Developer",
                    "period": "Jan 2024 - Present",
                    "bullets": [{"text": "Built and maintained a payments dashboard."}],
                }
            ],
            "skills": [{"category": "Core Stack", "items": ["Angular", "TypeScript"]}],
        },
        "leftovers": ["I hereby give consent for my personal data to be processed."],
    }

    profile = parse_resume_text(JANE_TEXT, llm=fake_llm, source_upload_id="upload-1")

    assert profile.core.identity.full_name == "Jane Doe"
    # cv_filename_prefix wasn't in the model's response — derived deterministically.
    assert profile.core.identity.cv_filename_prefix == "Jane_Doe_CV"
    assert profile.core.roles[0].company == "Acme Corp"
    assert profile.core.roles[0].origin == "parsed"
    assert profile.core.skills[0].items == ["Angular", "TypeScript"]
    assert (
        profile.leftovers[0].text == "I hereby give consent for my personal data to be processed."
    )
    assert profile.leftovers[0].source_upload_id == "upload-1"
    # The model left contact empty — the deterministic pre-fill still runs.
    assert "jane.doe@example.com" in profile.core.identity.contact


def test_parse_with_llm_garbage_response_falls_back(fake_llm) -> None:
    fake_llm.judge_response = {"nonsense": True}

    profile = parse_resume_text(JANE_TEXT, llm=fake_llm)

    assert profile.core.roles == []
    assert profile.leftovers[0].text == JANE_TEXT.strip()


def test_parse_with_llm_exception_falls_back() -> None:
    def _raising_llm(**kwargs: object) -> dict:
        raise RuntimeError("boom")

    profile = parse_resume_text(JANE_TEXT, llm=_raising_llm)

    assert profile.core.roles == []
    assert profile.leftovers[0].text == JANE_TEXT.strip()


def test_parse_with_llm_failed_validation_falls_back(fake_llm) -> None:
    # No full_name anywhere -> validate() rejects it even though a role parsed fine.
    fake_llm.judge_response = {
        "core": {"roles": [{"company": "Acme Corp", "title": "Dev", "period": "2024"}]},
    }

    profile = parse_resume_text(JANE_TEXT, llm=fake_llm)

    assert profile.core.roles == []
    assert profile.leftovers[0].text == JANE_TEXT.strip()


def test_parse_with_llm_blank_text_never_calls_llm() -> None:
    calls: list[dict] = []

    def _llm(**kwargs: object) -> dict:
        calls.append(kwargs)
        return {}

    profile = parse_resume_text("   ", llm=_llm)

    assert calls == []
    assert profile.leftovers == []


def test_parse_with_llm_round_trips(fake_llm) -> None:
    fake_llm.judge_response = {
        "core": {
            "identity": {"full_name": "Jane Doe", "contact": "jane.doe@example.com"},
            "roles": [{"company": "Acme Corp", "title": "Dev", "period": "2024"}],
        },
    }

    profile = parse_resume_text(JANE_TEXT, llm=fake_llm)
    rebuilt = from_dict(to_dict(profile))

    assert rebuilt.core.identity.full_name == profile.core.identity.full_name
    assert rebuilt.core.roles[0].company == profile.core.roles[0].company
