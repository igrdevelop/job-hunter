"""Tests for the GDPR/RODO consent clause appended to generated CVs."""

import importlib

import pytest
from docx import Document

import generate_docs


def _minimal_resume() -> dict:
    return {
        "summary": "Senior Frontend Developer with 10+ years.",
        "skills": {
            "frontend": "Angular (2-21), TypeScript",
            "tools": "Jest, Git",
            "methodologies": "Agile",
            "languages": "English (Fluent)",
        },
        "experience": [
            {
                "title": "Senior Frontend Developer (Angular)",
                "company": "Fairmarkit",
                "period": "Jun 2025 - Mar 2026",
                "subtitle": "Procurement Platform",
                "bullets": ["Built things."],
                "stack_line": "Stack: Angular 21.",
            }
        ],
        "education": "BSTU - Bachelor",
        "courses": "Angular Advanced Course",
    }


def _render(lang: str, mode: str, monkeypatch) -> list[str]:
    """Build a resume DOCX and return its paragraph texts."""
    import hunter.config as config

    monkeypatch.setattr(config, "CV_GDPR_CLAUSE", mode, raising=False)
    doc = Document()
    generate_docs.build_resume(doc, _minimal_resume(), "Angular", lang)
    return [p.text for p in doc.paragraphs]


def test_en_cv_has_english_gdpr_clause(monkeypatch) -> None:
    texts = _render("EN", "both", monkeypatch)
    assert texts[-1] == generate_docs.GDPR_CLAUSE_EN
    assert "GDPR" in texts[-1]
    assert "RODO" not in texts[-1]


def test_pl_cv_has_polish_rodo_clause(monkeypatch) -> None:
    texts = _render("PL", "both", monkeypatch)
    assert texts[-1] == generate_docs.GDPR_CLAUSE_PL
    assert "RODO" in texts[-1]
    assert "ustawą z dnia 10 maja 2018" in texts[-1]


def test_mode_none_omits_clause(monkeypatch) -> None:
    en = _render("EN", "none", monkeypatch)
    pl = _render("PL", "none", monkeypatch)
    assert generate_docs.GDPR_CLAUSE_EN not in en
    assert generate_docs.GDPR_CLAUSE_PL not in pl


def test_mode_pl_only_skips_english_cv(monkeypatch) -> None:
    en = _render("EN", "pl", monkeypatch)
    pl = _render("PL", "pl", monkeypatch)
    assert generate_docs.GDPR_CLAUSE_EN not in en
    assert pl[-1] == generate_docs.GDPR_CLAUSE_PL


def test_clause_is_small_italic_grey(monkeypatch) -> None:
    doc = Document()
    import hunter.config as config

    monkeypatch.setattr(config, "CV_GDPR_CLAUSE", "both", raising=False)
    generate_docs.build_resume(doc, _minimal_resume(), "Angular", "PL")
    last = doc.paragraphs[-1]
    run = last.runs[0]
    assert run.font.italic is True
    assert run.font.size.pt == pytest.approx(7.5)
